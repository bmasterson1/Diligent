"""Tests for diff_excel and diff_docx helpers.

Tests: Excel structured diff (sheets, cells, rows, named ranges),
Word paragraph-level diff, lazy imports, fixture validity.
"""

import sys

import pytest


# ---- Fixtures: create Excel and Word files programmatically ----


@pytest.fixture
def excel_files(tmp_path):
    """Create two Excel files with known differences using openpyxl.

    sample_a.xlsx: 2 sheets ("Summary", "Detail"), 10 rows each, 3 named ranges.
    sample_b.xlsx: same 2 sheets + 1 new ("Notes"), "Detail" has 2 changed cells
                   and 1 added row, 1 new named range.
    """
    from openpyxl import Workbook
    from openpyxl.workbook.defined_name import DefinedName

    # --- File A ---
    wb_a = Workbook()
    ws_summary = wb_a.active
    ws_summary.title = "Summary"
    for row in range(1, 11):
        ws_summary.cell(row=row, column=1, value=f"summary-{row}")
        ws_summary.cell(row=row, column=2, value=row * 100)

    ws_detail = wb_a.create_sheet("Detail")
    for row in range(1, 11):
        ws_detail.cell(row=row, column=1, value=f"detail-{row}")
        ws_detail.cell(row=row, column=2, value=row * 10)
        ws_detail.cell(row=row, column=3, value=f"note-{row}")

    # Named ranges in file A
    dn1 = DefinedName("TotalRevenue", attr_text="Summary!$B$1:$B$10")
    wb_a.defined_names.add(dn1)
    dn2 = DefinedName("DetailRange", attr_text="Detail!$A$1:$C$10")
    wb_a.defined_names.add(dn2)
    dn3 = DefinedName("SummaryHeader", attr_text="Summary!$A$1:$B$1")
    wb_a.defined_names.add(dn3)

    path_a = tmp_path / "sample_a.xlsx"
    wb_a.save(str(path_a))

    # --- File B ---
    wb_b = Workbook()
    ws_summary_b = wb_b.active
    ws_summary_b.title = "Summary"
    for row in range(1, 11):
        ws_summary_b.cell(row=row, column=1, value=f"summary-{row}")
        ws_summary_b.cell(row=row, column=2, value=row * 100)

    ws_detail_b = wb_b.create_sheet("Detail")
    for row in range(1, 11):
        ws_detail_b.cell(row=row, column=1, value=f"detail-{row}")
        ws_detail_b.cell(row=row, column=2, value=row * 10)
        ws_detail_b.cell(row=row, column=3, value=f"note-{row}")
    # Change 2 cells in Detail sheet
    ws_detail_b.cell(row=3, column=2, value=999)    # was 30
    ws_detail_b.cell(row=7, column=3, value="UPDATED")  # was "note-7"
    # Add 1 extra row
    ws_detail_b.cell(row=11, column=1, value="detail-11")
    ws_detail_b.cell(row=11, column=2, value=110)
    ws_detail_b.cell(row=11, column=3, value="note-11")

    # New sheet in B
    wb_b.create_sheet("Notes")

    # Named ranges in file B: keep same 3, add 1 new
    dn1b = DefinedName("TotalRevenue", attr_text="Summary!$B$1:$B$10")
    wb_b.defined_names.add(dn1b)
    dn2b = DefinedName("DetailRange", attr_text="Detail!$A$1:$C$10")
    wb_b.defined_names.add(dn2b)
    dn3b = DefinedName("SummaryHeader", attr_text="Summary!$A$1:$B$1")
    wb_b.defined_names.add(dn3b)
    dn4b = DefinedName("NotesRange", attr_text="Notes!$A$1:$A$10")
    wb_b.defined_names.add(dn4b)

    path_b = tmp_path / "sample_b.xlsx"
    wb_b.save(str(path_b))

    return str(path_a), str(path_b)


@pytest.fixture
def docx_files(tmp_path):
    """Create two Word files with known differences using python-docx.

    sample_a.docx: 5 paragraphs of text.
    sample_b.docx: 6 paragraphs, 2 changed from sample_a.
    """
    from docx import Document

    # --- File A ---
    doc_a = Document()
    paragraphs_a = [
        "Revenue grew 15% year over year.",
        "Customer retention rate was 92%.",
        "Operating expenses increased by $2M.",
        "Net income margin improved to 18%.",
        "Cash reserves stand at $45M.",
    ]
    for p in paragraphs_a:
        doc_a.add_paragraph(p)
    path_a = tmp_path / "sample_a.docx"
    doc_a.save(str(path_a))

    # --- File B ---
    doc_b = Document()
    paragraphs_b = [
        "Revenue grew 22% year over year.",      # changed (was 15%)
        "Customer retention rate was 92%.",       # unchanged
        "Operating expenses increased by $3.5M.", # changed (was $2M)
        "Net income margin improved to 18%.",     # unchanged
        "Cash reserves stand at $45M.",           # unchanged
        "New product line launched in Q4.",        # added
    ]
    for p in paragraphs_b:
        doc_b.add_paragraph(p)
    path_b = tmp_path / "sample_b.docx"
    doc_b.save(str(path_b))

    return str(path_a), str(path_b)


# ---- Excel diff tests ----


class TestDiffExcelSummary:
    def test_returns_dict_with_expected_keys(self, excel_files):
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(*excel_files)
        expected_keys = {
            "sheets_changed",
            "total_sheets",
            "cells_differ",
            "rows_added",
            "rows_removed",
            "named_ranges_added",
            "named_ranges_removed",
            "changed_sheet_names",
        }
        assert set(result.keys()) == expected_keys

    def test_detects_cell_value_differences(self, excel_files):
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(*excel_files)
        # 2 cells changed in Detail sheet
        assert result["cells_differ"] >= 2

    def test_detects_added_sheet(self, excel_files):
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(*excel_files)
        # File B has "Notes" sheet not in A
        assert result["total_sheets"] == 3  # Summary, Detail, Notes
        assert result["sheets_changed"] >= 1  # Detail has cell diffs

    def test_detects_row_count_differences(self, excel_files):
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(*excel_files)
        # Detail sheet in B has 1 extra row
        assert result["rows_added"] >= 1

    def test_detects_named_range_differences(self, excel_files):
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(*excel_files)
        # B has 1 new named range (NotesRange)
        assert result["named_ranges_added"] == 1
        assert result["named_ranges_removed"] == 0

    def test_changed_sheet_names_includes_detail(self, excel_files):
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(*excel_files)
        assert "Detail" in result["changed_sheet_names"]

    def test_identical_files_show_no_differences(self, excel_files):
        path_a, _ = excel_files
        from diligent.helpers.diff_excel import diff_excel_summary

        result = diff_excel_summary(path_a, path_a)
        assert result["cells_differ"] == 0
        assert result["rows_added"] == 0
        assert result["rows_removed"] == 0
        assert result["named_ranges_added"] == 0
        assert result["named_ranges_removed"] == 0
        assert result["changed_sheet_names"] == []


class TestDiffExcelLazyImport:
    def test_openpyxl_not_imported_at_module_level(self):
        """openpyxl should only be imported inside the function body."""
        # Remove openpyxl from sys.modules if present
        modules_before = set(sys.modules.keys())

        # Import the module (not the function call)
        import importlib
        if "diligent.helpers.diff_excel" in sys.modules:
            importlib.reload(sys.modules["diligent.helpers.diff_excel"])
        else:
            import diligent.helpers.diff_excel  # noqa: F401

        # Check no openpyxl was loaded by the import
        new_modules = set(sys.modules.keys()) - modules_before
        openpyxl_modules = [m for m in new_modules if m.startswith("openpyxl")]
        assert openpyxl_modules == [], (
            f"openpyxl was imported at module level: {openpyxl_modules}"
        )


# ---- Word diff tests ----


class TestDiffDocxSummary:
    def test_returns_dict_with_expected_keys(self, docx_files):
        from diligent.helpers.diff_docx import diff_docx_summary

        result = diff_docx_summary(*docx_files)
        expected_keys = {
            "paragraphs_changed",
            "paragraphs_added",
            "paragraphs_removed",
            "total_paragraphs_a",
            "total_paragraphs_b",
        }
        assert set(result.keys()) == expected_keys

    def test_detects_paragraph_level_text_differences(self, docx_files):
        from diligent.helpers.diff_docx import diff_docx_summary

        result = diff_docx_summary(*docx_files)
        # 2 changed paragraphs + 1 added
        assert result["paragraphs_changed"] >= 2
        assert result["paragraphs_added"] >= 1
        assert result["total_paragraphs_a"] == 5
        assert result["total_paragraphs_b"] == 6

    def test_identical_files_show_no_differences(self, docx_files):
        path_a, _ = docx_files
        from diligent.helpers.diff_docx import diff_docx_summary

        result = diff_docx_summary(path_a, path_a)
        assert result["paragraphs_changed"] == 0
        assert result["paragraphs_added"] == 0
        assert result["paragraphs_removed"] == 0


class TestDiffDocxVerbose:
    def test_verbose_returns_list_of_diff_lines(self, docx_files):
        from diligent.helpers.diff_docx import diff_docx_verbose

        result = diff_docx_verbose(*docx_files)
        assert isinstance(result, list)
        assert len(result) > 0
        # Should contain unified diff markers
        has_plus = any(line.startswith("+") for line in result)
        has_minus = any(line.startswith("-") for line in result)
        assert has_plus or has_minus


class TestDiffDocxLazyImport:
    def test_python_docx_not_imported_at_module_level(self):
        """python-docx should only be imported inside the function body."""
        modules_before = set(sys.modules.keys())

        import importlib
        if "diligent.helpers.diff_docx" in sys.modules:
            importlib.reload(sys.modules["diligent.helpers.diff_docx"])
        else:
            import diligent.helpers.diff_docx  # noqa: F401

        new_modules = set(sys.modules.keys()) - modules_before
        docx_modules = [m for m in new_modules if m.startswith("docx")]
        assert docx_modules == [], (
            f"python-docx was imported at module level: {docx_modules}"
        )


# ---- Fixture validity tests ----


class TestFixtureValidity:
    def test_excel_fixtures_are_valid_files(self, excel_files):
        from openpyxl import load_workbook

        for path in excel_files:
            wb = load_workbook(path, read_only=True)
            assert len(wb.sheetnames) >= 2
            wb.close()

    def test_docx_fixtures_are_valid_files(self, docx_files):
        from docx import Document

        for path in docx_files:
            doc = Document(path)
            assert len(doc.paragraphs) >= 5
