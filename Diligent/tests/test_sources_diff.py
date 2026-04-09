"""Tests for sources diff command and ingest auto-diff integration.

Tests: sources diff with Excel/Word/unsupported formats, --json and --verbose,
unknown source IDs, path resolution from SOURCES.md. Also tests ingest
auto-diff on --supersedes for Excel files.
"""

import json
import os
from pathlib import Path

import pytest
from click.testing import CliRunner

from diligent.cli import cli


@pytest.fixture
def runner():
    return CliRunner()


@pytest.fixture
def deal_dir(tmp_path, monkeypatch):
    """Create a deal folder with .diligence/, config, and empty SOURCES.md."""
    diligence = tmp_path / ".diligence"
    diligence.mkdir()

    config = {
        "schema_version": 1,
        "deal_code": "ARRIVAL",
        "created": "2026-01-01T00:00:00Z",
        "anchor_tolerance_pct": 1.0,
        "recent_window_days": 7,
        "workstreams": ["financial", "legal"],
    }
    (diligence / "config.json").write_text(json.dumps(config, indent=2), encoding="utf-8")
    (diligence / "SOURCES.md").write_text("# Sources\n", encoding="utf-8")

    monkeypatch.chdir(tmp_path)
    return tmp_path


@pytest.fixture
def excel_pair(deal_dir):
    """Create two Excel files with known differences and register them in SOURCES.md."""
    from openpyxl import Workbook
    from openpyxl.workbook.defined_name import DefinedName

    docs_dir = deal_dir / "documents"
    docs_dir.mkdir()

    # File A
    wb_a = Workbook()
    ws = wb_a.active
    ws.title = "Summary"
    for row in range(1, 6):
        ws.cell(row=row, column=1, value=f"item-{row}")
        ws.cell(row=row, column=2, value=row * 100)
    dn1 = DefinedName("TotalRev", attr_text="Summary!$B$1:$B$5")
    wb_a.defined_names.add(dn1)
    path_a = docs_dir / "report_v1.xlsx"
    wb_a.save(str(path_a))

    # File B
    wb_b = Workbook()
    ws_b = wb_b.active
    ws_b.title = "Summary"
    for row in range(1, 6):
        ws_b.cell(row=row, column=1, value=f"item-{row}")
        ws_b.cell(row=row, column=2, value=row * 100)
    ws_b.cell(row=2, column=2, value=999)  # changed cell
    ws_b.cell(row=6, column=1, value="item-6")  # added row
    ws_b.cell(row=6, column=2, value=600)
    dn1b = DefinedName("TotalRev", attr_text="Summary!$B$1:$B$6")
    wb_b.defined_names.add(dn1b)
    dn2b = DefinedName("NewRange", attr_text="Summary!$A$1:$A$6")
    wb_b.defined_names.add(dn2b)
    path_b = docs_dir / "report_v2.xlsx"
    wb_b.save(str(path_b))

    # Register both in SOURCES.md
    from diligent.state.models import SourceEntry, SourcesFile
    from diligent.state.sources import write_sources

    sources = SourcesFile(sources=[
        SourceEntry(
            id="ARRIVAL-001",
            path="documents/report_v1.xlsx",
            date_received="2026-03-01",
        ),
        SourceEntry(
            id="ARRIVAL-002",
            path="documents/report_v2.xlsx",
            date_received="2026-04-01",
            supersedes="ARRIVAL-001",
        ),
    ])
    write_sources(deal_dir / ".diligence" / "SOURCES.md", sources)

    return path_a, path_b


@pytest.fixture
def docx_pair(deal_dir):
    """Create two Word files with known differences and register them in SOURCES.md."""
    from docx import Document

    docs_dir = deal_dir / "documents"
    docs_dir.mkdir(exist_ok=True)

    # File A
    doc_a = Document()
    doc_a.add_paragraph("Revenue grew 15%.")
    doc_a.add_paragraph("Expenses were stable.")
    doc_a.add_paragraph("Net income positive.")
    path_a = docs_dir / "memo_v1.docx"
    doc_a.save(str(path_a))

    # File B
    doc_b = Document()
    doc_b.add_paragraph("Revenue grew 22%.")  # changed
    doc_b.add_paragraph("Expenses were stable.")
    doc_b.add_paragraph("Net income positive.")
    doc_b.add_paragraph("New product launched.")  # added
    path_b = docs_dir / "memo_v2.docx"
    doc_b.save(str(path_b))

    # Register both
    from diligent.state.models import SourceEntry, SourcesFile
    from diligent.state.sources import write_sources

    sources = SourcesFile(sources=[
        SourceEntry(
            id="ARRIVAL-001",
            path="documents/memo_v1.docx",
            date_received="2026-03-01",
        ),
        SourceEntry(
            id="ARRIVAL-002",
            path="documents/memo_v2.docx",
            date_received="2026-04-01",
            supersedes="ARRIVAL-001",
        ),
    ])
    write_sources(deal_dir / ".diligence" / "SOURCES.md", sources)

    return path_a, path_b


@pytest.fixture
def unsupported_pair(deal_dir):
    """Create two PDF files and register them in SOURCES.md."""
    docs_dir = deal_dir / "documents"
    docs_dir.mkdir(exist_ok=True)

    path_a = docs_dir / "contract_v1.pdf"
    path_a.write_text("fake pdf a", encoding="utf-8")
    path_b = docs_dir / "contract_v2.pdf"
    path_b.write_text("fake pdf b", encoding="utf-8")

    from diligent.state.models import SourceEntry, SourcesFile
    from diligent.state.sources import write_sources

    sources = SourcesFile(sources=[
        SourceEntry(
            id="ARRIVAL-001",
            path="documents/contract_v1.pdf",
            date_received="2026-03-01",
        ),
        SourceEntry(
            id="ARRIVAL-002",
            path="documents/contract_v2.pdf",
            date_received="2026-04-01",
        ),
    ])
    write_sources(deal_dir / ".diligence" / "SOURCES.md", sources)

    return path_a, path_b


# ---- Sources diff command tests ----


class TestSourcesDiffExcel:
    def test_diff_excel_shows_structured_summary(self, runner, excel_pair):
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-002"],
                               catch_exceptions=False)
        assert result.exit_code == 0, result.output
        assert "Diff: ARRIVAL-001 vs ARRIVAL-002" in result.output
        assert "sheets changed:" in result.output
        assert "cells differ:" in result.output
        assert "rows added:" in result.output
        assert "rows removed:" in result.output
        assert "named ranges:" in result.output

    def test_diff_excel_json_output(self, runner, excel_pair):
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-002", "--json"],
                               catch_exceptions=False)
        assert result.exit_code == 0, result.output
        data = json.loads(result.output)
        assert data["diff_type"] == "excel"
        assert "cells_differ" in data
        assert "sheets_changed" in data


class TestSourcesDiffDocx:
    def test_diff_docx_shows_paragraph_summary(self, runner, docx_pair):
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-002"],
                               catch_exceptions=False)
        assert result.exit_code == 0, result.output
        assert "Diff: ARRIVAL-001 vs ARRIVAL-002" in result.output
        assert "paragraphs changed:" in result.output
        assert "paragraphs added:" in result.output
        assert "paragraphs removed:" in result.output

    def test_diff_docx_verbose_shows_actual_diffs(self, runner, docx_pair):
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-002", "--verbose"],
                               catch_exceptions=False)
        assert result.exit_code == 0, result.output
        # Verbose output should have unified diff markers
        assert "---" in result.output or "+++" in result.output


class TestSourcesDiffUnsupported:
    def test_diff_unsupported_format_shows_message(self, runner, unsupported_pair):
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-002"],
                               catch_exceptions=False)
        assert result.exit_code == 0, result.output
        assert "Diff not supported for" in result.output


class TestSourcesDiffErrors:
    def test_unknown_source_id_exits_1(self, runner, deal_dir):
        # Empty SOURCES.md, no entries
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-999", "ARRIVAL-998"],
                               catch_exceptions=False)
        assert result.exit_code != 0

    def test_one_unknown_id_exits_1(self, runner, excel_pair):
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-999"],
                               catch_exceptions=False)
        assert result.exit_code != 0


class TestSourcesDiffPathResolution:
    def test_resolves_paths_from_sources_md_relative_to_deal_root(self, runner, excel_pair):
        """Paths in SOURCES.md are relative to deal root (parent of .diligence/)."""
        result = runner.invoke(cli, ["sources", "diff", "ARRIVAL-001", "ARRIVAL-002"],
                               catch_exceptions=False)
        assert result.exit_code == 0, result.output
        # Should have successfully loaded the files (no file not found error)
        assert "Diff: ARRIVAL-001 vs ARRIVAL-002" in result.output


# ---- Ingest auto-diff tests ----


class TestIngestAutoDiff:
    def test_ingest_with_supersedes_excel_shows_diff(self, runner, deal_dir):
        """Ingesting with --supersedes pointing to an Excel source auto-prints diff."""
        from openpyxl import Workbook

        docs = deal_dir / "documents"
        docs.mkdir(exist_ok=True)

        # Create and ingest first Excel file
        wb1 = Workbook()
        ws1 = wb1.active
        ws1.title = "Data"
        for r in range(1, 6):
            ws1.cell(row=r, column=1, value=f"row-{r}")
        wb1.save(str(docs / "data_v1.xlsx"))

        result1 = runner.invoke(cli, [
            "ingest", str(docs / "data_v1.xlsx"),
            "--date", "2026-03-01",
        ], catch_exceptions=False)
        assert result1.exit_code == 0, result1.output

        # Create second Excel file with differences
        wb2 = Workbook()
        ws2 = wb2.active
        ws2.title = "Data"
        for r in range(1, 7):  # 1 extra row
            ws2.cell(row=r, column=1, value=f"row-{r}")
        ws2.cell(row=1, column=1, value="CHANGED")  # changed cell
        wb2.save(str(docs / "data_v2.xlsx"))

        # Ingest with --supersedes
        result2 = runner.invoke(cli, [
            "ingest", str(docs / "data_v2.xlsx"),
            "--date", "2026-04-01",
            "--supersedes", "ARRIVAL-001",
        ], catch_exceptions=False)
        assert result2.exit_code == 0, result2.output
        # Should show compact diff inline
        assert "Diff vs ARRIVAL-001" in result2.output
        assert "sheets changed:" in result2.output or "cells differ:" in result2.output

    def test_ingest_auto_diff_compact_format(self, runner, deal_dir):
        """Auto-diff format matches locked format from CONTEXT.md."""
        from openpyxl import Workbook

        docs = deal_dir / "documents"
        docs.mkdir(exist_ok=True)

        # First file
        wb1 = Workbook()
        ws1 = wb1.active
        ws1.title = "Sheet1"
        ws1.cell(row=1, column=1, value="A")
        wb1.save(str(docs / "file_a.xlsx"))

        runner.invoke(cli, ["ingest", str(docs / "file_a.xlsx")], catch_exceptions=False)

        # Second file
        wb2 = Workbook()
        ws2 = wb2.active
        ws2.title = "Sheet1"
        ws2.cell(row=1, column=1, value="B")  # different
        wb2.save(str(docs / "file_b.xlsx"))

        result = runner.invoke(cli, [
            "ingest", str(docs / "file_b.xlsx"),
            "--supersedes", "ARRIVAL-001",
        ], catch_exceptions=False)
        assert result.exit_code == 0, result.output
        # Must have the "Run `diligent sources diff ..." hint
        assert "diligent sources diff" in result.output

    def test_ingest_with_supersedes_non_excel_skips_diff(self, runner, deal_dir):
        """Ingest with --supersedes pointing to non-Excel source skips auto-diff."""
        docs = deal_dir / "documents"
        docs.mkdir(exist_ok=True)

        # Ingest a PDF first
        pdf = docs / "report.pdf"
        pdf.write_text("fake pdf", encoding="utf-8")
        result1 = runner.invoke(cli, ["ingest", str(pdf)], catch_exceptions=False)
        assert result1.exit_code == 0

        # Ingest another PDF with --supersedes
        pdf2 = docs / "report_v2.pdf"
        pdf2.write_text("fake pdf v2", encoding="utf-8")
        result2 = runner.invoke(cli, [
            "ingest", str(pdf2),
            "--supersedes", "ARRIVAL-001",
        ], catch_exceptions=False)
        assert result2.exit_code == 0, result2.output
        # No diff output for PDFs
        assert "Diff vs" not in result2.output
        assert "sheets changed:" not in result2.output

    def test_ingest_auto_diff_failure_does_not_block_ingest(self, runner, deal_dir):
        """If auto-diff fails (e.g., missing file), ingest still succeeds."""
        from openpyxl import Workbook

        docs = deal_dir / "documents"
        docs.mkdir(exist_ok=True)

        # Register a source entry manually that points to a missing file
        from diligent.state.models import SourceEntry, SourcesFile
        from diligent.state.sources import write_sources

        sources = SourcesFile(sources=[
            SourceEntry(
                id="ARRIVAL-001",
                path="documents/missing.xlsx",
                date_received="2026-03-01",
            ),
        ])
        write_sources(deal_dir / ".diligence" / "SOURCES.md", sources)

        # Create a new Excel file
        wb = Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="data")
        wb.save(str(docs / "new.xlsx"))

        # Ingest with --supersedes pointing to the missing file
        result = runner.invoke(cli, [
            "ingest", str(docs / "new.xlsx"),
            "--supersedes", "ARRIVAL-001",
        ], catch_exceptions=False)
        # Ingest should succeed even if diff fails
        assert result.exit_code == 0, result.output
        assert "Registered ARRIVAL-002" in result.output
